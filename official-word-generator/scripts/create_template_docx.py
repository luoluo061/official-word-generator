from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ensure_dependencies import ensure as ensure_dependencies  # noqa: E402
from list_features import parse_feature_catalog  # noqa: E402
from profile_resolver import ProfileError, resolve_profile  # noqa: E402


ensure_dependencies()

from docx import Document  # noqa: E402
from docx.enum.style import WD_STYLE_TYPE  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402


SKILL_DIR = SCRIPT_DIR.parent
PROFILES_DIR = SKILL_DIR / "profiles"
BASE_TEMPLATE = SKILL_DIR / "assets" / "base-official-template.docx"
BLACK = RGBColor(0, 0, 0)
FONT_BODY = "仿宋_GB2312"
FONT_TITLE = "方正小标宋简体"
FONT_HEI = "黑体"
FONT_KAI = "楷体_GB2312"
FONT_ASCII = "Times New Roman"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def feature_status_map() -> dict[str, str]:
    features, _warnings = parse_feature_catalog()
    return {feature.feature_id: feature.status for feature in features}


def enabled_feature_ids(features: dict) -> list[str]:
    enabled = set(features.get("enabled_features") or [])
    for feature_id, config in (features.get("features") or {}).items():
        if isinstance(config, dict) and config.get("enabled"):
            enabled.add(feature_id)
    return sorted(enabled)


def set_first_line_chars(style, chars: int) -> None:
    element = getattr(style, "element", None)
    ppr = element.get_or_add_pPr() if element is not None else style._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    for attr in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
        ind.attrib.pop(qn(attr), None)
    if chars:
        ind.set(qn("w:firstLineChars"), str(chars * 100))


def ensure_paragraph_style(doc: Document, name: str, font: str, size_pt: float, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_chars: int = 0, line_rule=WD_LINE_SPACING.EXACTLY, line_pt: float = 28, outline_level: int | None = None) -> None:
    style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = FONT_ASCII
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = line_rule
    fmt.line_spacing = 1.0 if line_rule == WD_LINE_SPACING.SINGLE else Pt(line_pt)
    set_first_line_chars(style, first_line_chars)
    if outline_level is not None:
        ppr = style.element.get_or_add_pPr()
        for child in ppr.findall(qn("w:outlineLvl")):
            ppr.remove(child)
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(outline_level))
        ppr.append(outline)


def ensure_candidate_styles(doc: Document) -> list[str]:
    generated = []
    specs = [
        ("公文标题", FONT_TITLE, 22, False, WD_ALIGN_PARAGRAPH.CENTER, 0, WD_LINE_SPACING.AT_LEAST, 0, None),
        ("一级标题", FONT_HEI, 16, False, WD_ALIGN_PARAGRAPH.LEFT, 2, WD_LINE_SPACING.EXACTLY, 28, 0),
        ("二级标题", FONT_KAI, 16, False, WD_ALIGN_PARAGRAPH.LEFT, 2, WD_LINE_SPACING.EXACTLY, 28, 1),
        ("三级标题", FONT_KAI, 16, False, WD_ALIGN_PARAGRAPH.LEFT, 2, WD_LINE_SPACING.EXACTLY, 28, 2),
        ("图题", FONT_BODY, 16, False, WD_ALIGN_PARAGRAPH.CENTER, 0, WD_LINE_SPACING.EXACTLY, 28, None),
        ("表题", FONT_BODY, 16, False, WD_ALIGN_PARAGRAPH.CENTER, 0, WD_LINE_SPACING.EXACTLY, 28, None),
        ("表格正文", FONT_BODY, 16, False, WD_ALIGN_PARAGRAPH.CENTER, 0, WD_LINE_SPACING.SINGLE, 0, None),
        ("表格表头", FONT_BODY, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 0, WD_LINE_SPACING.SINGLE, 0, None),
        ("备注", FONT_BODY, 16, False, WD_ALIGN_PARAGRAPH.LEFT, 0, WD_LINE_SPACING.EXACTLY, 28, None),
    ]
    for spec in specs:
        name, font, size_pt, bold, align, first_line_chars, line_rule, line_pt, outline_level = spec
        ensure_paragraph_style(
            doc,
            name,
            font,
            size_pt,
            bold=bold,
            align=align,
            first_line_chars=first_line_chars,
            line_rule=line_rule,
            line_pt=line_pt,
            outline_level=outline_level,
        )
        generated.append(spec[0])
    ensure_paragraph_style(doc, "Normal", FONT_BODY, 16, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_chars=2)
    generated.append("Normal")
    return generated


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_toc_field(doc: Document) -> None:
    doc.add_paragraph("目录", style="一级标题")
    paragraph = doc.add_paragraph(style="Normal")
    set_first_line_chars(paragraph, 0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "请在 Word/WPS 中更新域以生成目录。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_candidate_placeholders(doc: Document, feature_ids: set[str], skipped_planned: list[str]) -> list[str]:
    sections = []
    if "structure.cover" in feature_ids:
        skipped_planned.append("structure.cover")

    doc.add_paragraph("模板候选说明", style="公文标题")
    doc.add_paragraph("本文件为 Codex 根据 profile 功能配置自动生成的 Word 模板候选，仅用于 draft 测试，不代表 production 模板。", style="Normal")
    sections.append("candidate_notice")

    if "toc.insert_field" in feature_ids:
        add_page_break(doc)
        add_toc_field(doc)
        sections.append("toc_field")
        if "toc.page_breaks" in feature_ids:
            add_page_break(doc)

    doc.add_paragraph("一、一级标题示例", style="一级标题")
    doc.add_paragraph("（一）二级标题示例", style="二级标题")
    doc.add_paragraph("1. 三级标题示例", style="三级标题")
    doc.add_paragraph("正文样式示例。", style="Normal")
    sections.append("body_style_examples")

    if "table.markdown_to_word" in feature_ids:
        doc.add_paragraph("表1  表格样式示例", style="表题")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "表头"
        table.cell(0, 1).text = "说明"
        table.cell(1, 0).text = "正文"
        table.cell(1, 1).text = "内容"
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = "表格表头" if row_index == 0 else "表格正文"
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sections.append("table_example")

    if "caption.figure" in feature_ids or "image.insert_inline" in feature_ids:
        doc.add_paragraph("图1  图题样式示例", style="图题")
        sections.append("figure_caption_example")

    if "structure.signature_block" in feature_ids:
        skipped_planned.append("structure.signature_block")
    if "structure.appendix" in feature_ids:
        skipped_planned.append("structure.appendix")
    return sections


def render_report(profile_id: str, source_template: Path, output_template: Path, enabled: list[str], skipped_planned: list[str], partial_warnings: list[str], generated_styles: list[str], generated_sections: list[str]) -> str:
    lines = [
        "# Template Build Report",
        "",
        f"- profile_id: `{profile_id}`",
        f"- source_template: `{source_template}`",
        f"- output_template: `{output_template}`",
        "",
        "## Enabled Features",
        "",
    ]
    lines.extend([f"- `{feature_id}`" for feature_id in enabled] if enabled else ["- none"])
    lines.extend(["", "## Skipped Planned Features", ""])
    skipped_unique = sorted(set(skipped_planned))
    lines.extend([f"- `{feature_id}`" for feature_id in skipped_unique] if skipped_unique else ["- none"])
    lines.extend(["", "## Partial Feature Warnings", ""])
    lines.extend([f"- `{feature_id}`" for feature_id in partial_warnings] if partial_warnings else ["- none"])
    lines.extend(["", "## Generated Styles", ""])
    lines.extend([f"- `{style}`" for style in generated_styles] if generated_styles else ["- none"])
    lines.extend(["", "## Generated Sections", ""])
    lines.extend([f"- `{section}`" for section in generated_sections] if generated_sections else ["- none"])
    lines.extend(
        [
            "",
            "## Next Steps",
            "",
            "- Run `py scripts\\validate_profile.py --profile <profile_id>`.",
            "- Prepare or refine `example.md`.",
            "- Run `generate_docx.py --profile` and `validate_docx.py --profile`.",
            "- Open the generated Word file for visual review.",
            "- Keep this profile as draft until all validation and visual review pass.",
            "",
        ]
    )
    return "\n".join(lines)


def create_template(profile_id: str, source_profile: str, force: bool = False) -> Path:
    profile_info = resolve_profile(profile_id)
    source_info = resolve_profile(source_profile)
    if profile_info["status"] == "production" and not force:
        raise ValueError("Refusing to modify a production profile template without --force.")

    profile_dir = Path(profile_info["profile_dir"])
    output_template = profile_dir / "template.docx"
    if output_template.exists() and not force:
        raise FileExistsError(f"template.docx already exists: {output_template}; use --force to overwrite.")

    source_template = Path(source_info["template_path"])
    if not source_template.exists() or source_template.suffix.lower() != ".docx":
        source_template = BASE_TEMPLATE
    if not source_template.exists():
        raise FileNotFoundError(f"source template not found: {source_template}")

    shutil.copy2(source_template, output_template)
    doc = Document(str(output_template))
    generated_styles = ensure_candidate_styles(doc)
    clear_body(doc)

    features = read_json(profile_dir / "features.json")
    statuses = feature_status_map()
    enabled = enabled_feature_ids(features)
    feature_set = set(enabled)
    skipped_planned = [feature_id for feature_id in enabled if statuses.get(feature_id) == "planned"]
    partial_warnings = [feature_id for feature_id in enabled if statuses.get(feature_id) == "partial"]
    generated_sections = add_candidate_placeholders(doc, feature_set, skipped_planned)
    doc.save(output_template)

    features["template"] = "template.docx"
    write_json(profile_dir / "features.json", features)

    report = render_report(profile_id, source_template, output_template, enabled, skipped_planned, partial_warnings, generated_styles, generated_sections)
    (profile_dir / "template_build_report.md").write_text(report, encoding="utf-8")
    return output_template


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a draft template.docx candidate for a profile.")
    parser.add_argument("--profile", required=True, help="Target profile id.")
    parser.add_argument("--from", dest="source_profile", default="general_official", help="Reference profile id.")
    parser.add_argument("--force", action="store_true", help="Overwrite existing template.docx. Required for production profiles.")
    args = parser.parse_args()
    try:
        output = create_template(args.profile, args.source_profile, args.force)
    except (ProfileError, ValueError, FileExistsError, FileNotFoundError) as exc:
        parser.error(str(exc))
    print(f"Created template candidate: {output}")
    print(f"Build report: {output.with_name('template_build_report.md')}")
    print(f"Next: py scripts\\validate_profile.py --profile {args.profile}")


if __name__ == "__main__":
    main()
