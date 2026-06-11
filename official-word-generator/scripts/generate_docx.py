from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile, ZIP_DEFLATED

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ensure_dependencies import ensure as ensure_dependencies  # noqa: E402

ensure_dependencies()

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from markdown_parser import Block, parse_markdown  # noqa: E402
from text_io import read_text_safely  # noqa: E402
from validate_docx import validate_document  # noqa: E402


BLACK = RGBColor(0, 0, 0)
FONT_TITLE = "方正小标宋简体"
FONT_BODY = "仿宋_GB2312"
FONT_L2 = "楷体_GB2312"
FONT_HEI = "黑体"
FONT_ASCII = "Times New Roman"


STYLE_CANDIDATES = {
    "title": ["公文标题", "Heading 1"],
    "h1": ["一级标题", "Heading 2"],
    "h2": ["二级标题", "Heading 3"],
    "body": ["Normal"],
    "figure_caption": ["图题", "Caption"],
    "table_caption": ["表题", "Caption"],
    "table_body": ["表格正文"],
    "table_header": ["表格表头"],
    "note": ["备注", "Intense Quote", "Quote"],
}

ALIGNMENTS = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


DEFAULT_RULES = {
    "styles": {
        "title": {"style_name": "公文标题", "east_asia": FONT_TITLE, "ascii": FONT_ASCII, "size_pt": 22, "bold": False, "first_line_chars": 0, "align": "center", "line_rule": "at_least", "line_pt": 0},
        "body": {"style_name": "Normal", "east_asia": FONT_BODY, "ascii": FONT_ASCII, "size_pt": 16, "bold": False, "first_line_chars": 2, "align": "justify", "line_rule": "exact", "line_pt": 28},
        "h1": {"style_name": "一级标题", "east_asia": FONT_HEI, "ascii": FONT_ASCII, "size_pt": 16, "bold": True, "first_line_chars": 2, "align": "left", "line_rule": "exact", "line_pt": 28, "outline_level": 0},
        "h2": {"style_name": "二级标题", "east_asia": FONT_L2, "ascii": FONT_ASCII, "size_pt": 16, "bold": False, "first_line_chars": 2, "align": "left", "line_rule": "exact", "line_pt": 28, "outline_level": 1},
        "figure_caption": {"style_name": "图题", "east_asia": FONT_BODY, "ascii": FONT_ASCII, "size_pt": 16, "bold": False, "first_line_chars": 0, "align": "center", "line_rule": "exact", "line_pt": 28},
        "table_caption": {"style_name": "表题", "east_asia": FONT_BODY, "ascii": FONT_ASCII, "size_pt": 16, "bold": False, "first_line_chars": 0, "align": "center", "line_rule": "exact", "line_pt": 28},
        "table_body": {"style_name": "表格正文", "east_asia": FONT_BODY, "ascii": FONT_ASCII, "size_pt": 10.5, "bold": False, "first_line_chars": 0, "align": "center", "line_rule": "single", "line_pt": 0},
        "table_header": {"style_name": "表格表头", "east_asia": FONT_HEI, "ascii": FONT_ASCII, "size_pt": 10.5, "bold": True, "first_line_chars": 0, "align": "center", "line_rule": "single", "line_pt": 0},
        "note": {"style_name": "备注", "east_asia": FONT_BODY, "ascii": FONT_ASCII, "size_pt": 16, "bold": False, "first_line_chars": 0, "align": "left", "line_rule": "exact", "line_pt": 28},
    }
}


def load_format_rules(rules_path: Path | None = None) -> dict:
    if rules_path is None:
        rules_path = SCRIPT_DIR.parent / "references" / "format_rules.md"
    if not rules_path.exists():
        return DEFAULT_RULES
    text = read_text_safely(rules_path)
    match = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.DOTALL)
    if not match:
        return DEFAULT_RULES
    try:
        loaded = json.loads(match.group(1))
    except json.JSONDecodeError:
        return DEFAULT_RULES
    merged = json.loads(json.dumps(DEFAULT_RULES, ensure_ascii=False))
    for key, value in loaded.get("styles", {}).items():
        if key in merged["styles"] and isinstance(value, dict):
            merged["styles"][key].update(value)
    return merged


def set_run_font(run, east_asia: str, size_pt: float, *, bold: bool = False, ascii_font: str = FONT_ASCII) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def set_first_line_chars(paragraph_or_style, chars: int) -> None:
    """Use Word's character-based first-line indent instead of a point/cm approximation."""
    element = getattr(paragraph_or_style, "element", None)
    ppr = element.get_or_add_pPr() if element is not None else paragraph_or_style._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    for attr in ("w:firstLine", "w:firstLineChars", "w:hanging", "w:hangingChars"):
        attr_name = qn(attr)
        if attr_name in ind.attrib:
            del ind.attrib[attr_name]
    if chars:
        ind.set(qn("w:firstLineChars"), str(int(chars) * 100))


def configure_style(style, rule: dict) -> None:
    east_asia = rule.get("east_asia", FONT_BODY)
    ascii_font = rule.get("ascii", FONT_ASCII)
    size_pt = float(rule.get("size_pt", 16))
    style.font.name = ascii_font
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size_pt)
    style.font.bold = bool(rule.get("bold", False))
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    first_line_chars = int(rule.get("first_line_chars", 0))
    set_first_line_chars(style, first_line_chars)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    line_rule = rule.get("line_rule", "exact")
    line_pt = float(rule.get("line_pt", 28))
    if line_rule == "at_least":
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        fmt.line_spacing = Pt(line_pt)
    elif line_rule == "single":
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.line_spacing = 1.0
    else:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_pt)
    align = ALIGNMENTS.get(rule.get("align"))
    if align is not None:
        fmt.alignment = align
    if "outline_level" in rule:
        ppr = style.element.get_or_add_pPr()
        for child in ppr.findall(qn("w:outlineLvl")):
            ppr.remove(child)
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(rule["outline_level"]))
        ppr.append(outline)


def ensure_style(doc: Document, rule: dict):
    name = rule["style_name"]
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    configure_style(style, rule)
    return style


def ensure_fallback_styles(doc: Document, rules: dict) -> None:
    for rule in rules["styles"].values():
        if rule["style_name"] == "Normal":
            configure_style(doc.styles["Normal"], rule)
        else:
            ensure_style(doc, rule)


def resolve_style(doc: Document, key: str) -> str:
    for name in STYLE_CANDIDATES[key]:
        if name in doc.styles:
            return name
    return "Normal"


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_text_paragraph(doc: Document, text: str, style_name: str, *, align=None, first_line_chars: int | None = None):
    p = doc.add_paragraph(style=style_name)
    if align is not None:
        p.alignment = align
    if first_line_chars is not None:
        set_first_line_chars(p, first_line_chars)
    run = p.add_run(text)
    style = doc.styles[style_name]
    east_asia = style.element.rPr.rFonts.get(qn("w:eastAsia")) if style.element.rPr is not None else None
    set_run_font(run, east_asia or FONT_BODY, style.font.size.pt if style.font.size else 16, bold=bool(style.font.bold))
    return p


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table_of_contents(doc: Document) -> None:
    add_text_paragraph(doc, "目录", resolve_style(doc, "h1"), first_line_chars=2)
    p = doc.add_paragraph(style=resolve_style(doc, "body"))
    set_first_line_chars(p, 0)
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word/WPS 中更新域以生成目录。"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def update_fields_with_word(docx_path: Path) -> list[str]:
    warnings: list[str] = []
    try:
        ensure_dependencies(["win32com.client"])
        import win32com.client  # type: ignore
    except ImportError:
        return ["Word field update skipped: pywin32 is not installed."]

    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        doc.Save()
    except Exception as exc:  # pragma: no cover - depends on local Word installation
        warnings.append(f"Word field update failed: {exc}")
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
    return warnings


def normalize_after_word_field_update(docx_path: Path) -> None:
    # Word may expand TOC fields and rewrite official firstLineChars into twip-based firstLine.
    # Keep Word's TOC indentation, but restore non-TOC paragraphs/styles to character indent.
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = ns["w"]
    style_char_indents = {
        "Normal": 2,
        "公文标题": 0,
        "一级标题": 2,
        "二级标题": 2,
        "图题": 0,
        "表题": 0,
        "表格正文": 0,
        "表格表头": 0,
        "备注": 0,
        "附件标题": 0,
    }

    def attr(local_name: str) -> str:
        return f"{{{w_ns}}}{local_name}"

    def style_name(style_element) -> str | None:
        name = style_element.find("w:name", namespaces=ns)
        return name.get(attr("val")) if name is not None else None

    def remove_twip_indent(ind_element) -> None:
        for key in ("firstLine", "hanging"):
            ind_element.attrib.pop(attr(key), None)

    def set_chars(ind_element, chars: int) -> None:
        for key in ("firstLine", "hanging", "hangingChars"):
            ind_element.attrib.pop(attr(key), None)
        if chars:
            ind_element.set(attr("firstLineChars"), str(chars * 100))
        else:
            ind_element.attrib.pop(attr("firstLineChars"), None)

    with TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        with ZipFile(docx_path) as archive:
            archive.extractall(tmp_dir)

        styles_path = tmp_dir / "word" / "styles.xml"
        if styles_path.exists():
            tree = etree.parse(str(styles_path))
            for style in tree.xpath("//w:style", namespaces=ns):
                name = style_name(style)
                ind = style.find("w:pPr/w:ind", namespaces=ns)
                if ind is None:
                    continue
                if name and name.lower().startswith("toc"):
                    continue
                if name in style_char_indents:
                    set_chars(ind, style_char_indents[name])
                else:
                    remove_twip_indent(ind)
            tree.write(str(styles_path), encoding="utf-8", xml_declaration=True)

        document_path = tmp_dir / "word" / "document.xml"
        if document_path.exists():
            tree = etree.parse(str(document_path))
            for paragraph in tree.xpath("//w:p", namespaces=ns):
                p_style = paragraph.find("w:pPr/w:pStyle", namespaces=ns)
                style_val = p_style.get(attr("val")) if p_style is not None else ""
                if style_val.startswith("TOC"):
                    continue
                ind = paragraph.find("w:pPr/w:ind", namespaces=ns)
                if ind is None:
                    continue
                if attr("firstLine") in ind.attrib or attr("firstLineChars") in ind.attrib:
                    set_chars(ind, 2)
            tree.write(str(document_path), encoding="utf-8", xml_declaration=True)

        with ZipFile(docx_path, "w", ZIP_DEFLATED) as archive:
            for file_path in tmp_dir.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, file_path.relative_to(tmp_dir).as_posix())


def add_image(doc: Document, block: Block, content_dir: Path, missing_images: list[str]) -> None:
    image_path = Path(block.path or "")
    if not image_path.is_absolute():
        image_path = content_dir / image_path

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0

    if image_path.exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.5))
    else:
        missing_images.append(str(image_path))
        run = p.add_run(f"[Missing image: {image_path}]")
        set_run_font(run, FONT_BODY, 16)

    if block.caption:
        add_text_paragraph(doc, block.caption, resolve_style(doc, "figure_caption"), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_chars=0)


def add_table(doc: Document, block: Block) -> None:
    rows = block.rows or []
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_three_line_table_borders(table)
    set_table_width_pct(table, 100)
    for r_idx, row in enumerate(rows):
        if r_idx == 0:
            set_repeat_table_header(table.rows[r_idx])
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=60, start=80, bottom=60, end=80)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                paragraph.style = resolve_style(doc, "table_header" if r_idx == 0 else "table_body")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_first_line_chars(paragraph, 0)
                for run in paragraph.runs:
                    rule_key = "table_header" if r_idx == 0 else "table_body"
                    rule = load_format_rules()["styles"].get(rule_key, DEFAULT_RULES["styles"][rule_key])
                    set_run_font(
                        run,
                        rule.get("east_asia", FONT_BODY),
                        float(rule.get("size_pt", 10.5)),
                        bold=bool(rule.get("bold", False)),
                        ascii_font=rule.get("ascii", FONT_ASCII),
                    )


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_table_width_pct(table, pct: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct * 50))


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_three_line_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge, value, size in (
        ("top", "single", "12"),
        ("left", "nil", "0"),
        ("bottom", "single", "12"),
        ("right", "nil", "0"),
        ("insideH", "nil", "0"),
        ("insideV", "nil", "0"),
    ):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")
        borders.append(node)
    tbl_pr.append(borders)

    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            bottom = tc_borders.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                tc_borders.append(bottom)
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")


def generate_docx(template_path: Path, content_path: Path, output_path: Path, report_path: Path | None = None, rules_path: Path | None = None, update_fields: bool = False) -> dict:
    if template_path.suffix.lower() != ".docx":
        raise ValueError("Template must be .docx. Convert .doc to .docx before using this skill.")

    doc = Document(str(template_path))
    rules = load_format_rules(rules_path)
    ensure_fallback_styles(doc, rules)
    clear_body(doc)
    blocks = parse_markdown(content_path)
    missing_images: list[str] = []

    for block in blocks:
        if block.type == "heading":
            if block.level == 1:
                add_text_paragraph(doc, block.text, resolve_style(doc, "title"), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_chars=0)
            elif block.level == 2:
                add_text_paragraph(doc, block.text, resolve_style(doc, "h1"), first_line_chars=2)
            elif block.level == 3:
                add_text_paragraph(doc, block.text, resolve_style(doc, "h2"), first_line_chars=2)
            else:
                add_text_paragraph(doc, block.text, resolve_style(doc, "body"), first_line_chars=2)
        elif block.type == "paragraph":
            add_text_paragraph(doc, block.text, resolve_style(doc, "body"), first_line_chars=2)
        elif block.type == "quote":
            add_text_paragraph(doc, block.text, resolve_style(doc, "note"), first_line_chars=0)
        elif block.type == "toc":
            if any(paragraph.text.strip() for paragraph in doc.paragraphs):
                add_page_break(doc)
            add_table_of_contents(doc)
            add_page_break(doc)
        elif block.type == "image":
            add_image(doc, block, content_path.parent, missing_images)
        elif block.type == "table":
            add_table(doc, block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    field_warnings = update_fields_with_word(output_path) if update_fields else []
    if update_fields and not field_warnings:
        normalize_after_word_field_update(output_path)
    result = validate_document(output_path, report_path, content_path)
    if missing_images or field_warnings:
        result.setdefault("warnings", []).extend(f"Missing image: {path}" for path in missing_images)
        result.setdefault("warnings", []).extend(field_warnings)
        if report_path:
            from validate_docx import render_report

            report_path.write_text(render_report(result), encoding="utf-8")
    return result


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a Word docx from a template and Markdown content.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--content", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--rules", type=Path, help="Optional format_rules.md path. Defaults to the skill references file.")
    parser.add_argument("--update-fields", action="store_true", help="Use local Microsoft Word to update TOC/page fields after saving.")
    args = parser.parse_args()
    generate_docx(args.template, args.content, args.output, args.report, args.rules, args.update_fields)


if __name__ == "__main__":
    main()
