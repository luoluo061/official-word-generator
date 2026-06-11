from __future__ import annotations

import argparse
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ensure_dependencies import ensure as ensure_dependencies  # noqa: E402

ensure_dependencies()

from docx import Document


def _mm(value) -> float | None:
    if value is None:
        return None
    return round(value.mm, 2)


def inspect_template(template_path: Path, output_path: Path) -> None:
    doc = Document(str(template_path))
    section = doc.sections[0]

    lines = [
        "# Template Inspection",
        "",
        f"- Template: `{template_path}`",
        f"- Paragraphs: `{len(doc.paragraphs)}`",
        f"- Tables: `{len(doc.tables)}`",
        "",
        "## First Section Page Setup",
        "",
        f"- Page width: `{_mm(section.page_width)}` mm",
        f"- Page height: `{_mm(section.page_height)}` mm",
        f"- Top margin: `{_mm(section.top_margin)}` mm",
        f"- Bottom margin: `{_mm(section.bottom_margin)}` mm",
        f"- Left margin: `{_mm(section.left_margin)}` mm",
        f"- Right margin: `{_mm(section.right_margin)}` mm",
        f"- Header distance: `{_mm(section.header_distance)}` mm",
        f"- Footer distance: `{_mm(section.footer_distance)}` mm",
        "",
        "## Paragraph Styles",
        "",
    ]

    for style in doc.styles:
        if style.type != 1:
            continue
        font = style.font
        size = font.size.pt if font.size else None
        lines.append(f"- `{style.name}`: font=`{font.name}`, size=`{size}`, bold=`{font.bold}`")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect a Word template.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    inspect_template(args.template, args.output)


if __name__ == "__main__":
    main()
