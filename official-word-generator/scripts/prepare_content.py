from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ensure_dependencies import ensure as ensure_dependencies  # noqa: E402
from text_io import normalize_markdown_text, read_text_safely  # noqa: E402


def _looks_like_heading(text: str) -> bool:
    return bool(re.match(r"^[一二三四五六七八九十]+、", text)) or bool(re.match(r"^\d+(?:\.\d+)+\s+", text))


def _docx_to_markdown(input_path: Path) -> str:
    ensure_dependencies(["docx"])
    from docx import Document

    doc = Document(str(input_path))
    lines: list[str] = []
    title_written = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if not title_written:
            lines.extend([f"# {text}", ""])
            title_written = True
            continue
        if style_name in {"公文标题", "Title"}:
            lines.extend([f"# {text}", ""])
        elif style_name in {"一级标题", "Heading 2"} or re.match(r"^[一二三四五六七八九十]+、", text):
            lines.extend([f"## {text}", ""])
        elif style_name in {"二级标题", "Heading 3"} or re.match(r"^\d+(?:\.\d+)+\s+", text):
            lines.extend([f"### {text}", ""])
        elif _looks_like_heading(text) and len(text) <= 80:
            lines.extend([f"### {text}", ""])
        else:
            lines.extend([text, ""])

    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [cell.text.strip().replace("|", "｜").replace("\n", " ") for cell in row.cells]
            if any(values):
                rows.append(values)
        if not rows:
            continue
        col_count = max(len(row) for row in rows)
        normalized = [row + [""] * (col_count - len(row)) for row in rows]
        lines.append("| " + " | ".join(normalized[0]) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")
        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    return normalize_markdown_text("\n".join(lines))


def prepare_content(input_path: Path, output_path: Path) -> None:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        text = _docx_to_markdown(input_path)
    elif suffix in {".md", ".txt", ".text"}:
        text = normalize_markdown_text(read_text_safely(input_path))
    else:
        raise ValueError(f"Unsupported content input type: {input_path.suffix}. Use .md, .txt, or .docx.")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(text, encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize pasted text, Markdown, or docx content into UTF-8 Markdown.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    prepare_content(args.input, args.output)


if __name__ == "__main__":
    main()
