from __future__ import annotations

import argparse
import sys
from pathlib import Path
from zipfile import ZipFile

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from ensure_dependencies import ensure as ensure_dependencies  # noqa: E402

ensure_dependencies()

from docx import Document
from lxml import etree
from profile_resolver import ProfileError, resolve_profile  # noqa: E402


REQUIRED_STYLES = ["公文标题", "一级标题", "二级标题", "图题", "表题", "表格正文", "表格表头", "备注"]
EXPECTED_PAGE_SETUP_MM = {
    "page_width": 210,
    "page_height": 297,
    "top_margin": 37,
    "bottom_margin": 35,
    "left_margin": 27,
    "right_margin": 27,
    "header_distance": 15,
    "footer_distance": 28,
}
PAGE_TOLERANCE_MM = 0.35


def _mm(value) -> float | None:
    if value is None:
        return None
    return round(value.mm, 2)


def _xml_diagnostics(docx_path: Path) -> dict:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = ns["w"]

    def attr(local_name: str) -> str:
        return f"{{{w_ns}}}{local_name}"

    def style_name(style_element) -> str | None:
        name = style_element.find("w:name", namespaces=ns)
        return name.get(attr("val")) if name is not None else None

    with ZipFile(docx_path) as archive:
        names = archive.namelist()
        settings_xml = archive.read("word/settings.xml").decode("utf-8") if "word/settings.xml" in names else ""
        styles_xml = archive.read("word/styles.xml") if "word/styles.xml" in names else b""
        document_xml = archive.read("word/document.xml") if "word/document.xml" in names else b""
        styles_first_line_count = 0
        styles_first_line_chars_200_count = 0
        if styles_xml:
            styles_root = etree.fromstring(styles_xml)
            for style in styles_root.xpath("//w:style", namespaces=ns):
                name = style_name(style)
                if name and name.lower().startswith("toc"):
                    continue
                ind = style.find("w:pPr/w:ind", namespaces=ns)
                if ind is None:
                    continue
                if attr("firstLine") in ind.attrib:
                    styles_first_line_count += 1
                if ind.get(attr("firstLineChars")) == "200":
                    styles_first_line_chars_200_count += 1
        document_first_line_count = 0
        document_first_line_chars_200_count = 0
        table_count = 0
        tables_with_repeating_header = 0
        tables_with_three_line_borders = 0
        if document_xml:
            document_root = etree.fromstring(document_xml)
            for paragraph in document_root.xpath("//w:p", namespaces=ns):
                p_style = paragraph.find("w:pPr/w:pStyle", namespaces=ns)
                style_val = p_style.get(attr("val")) if p_style is not None else ""
                if style_val.startswith("TOC"):
                    continue
                ind = paragraph.find("w:pPr/w:ind", namespaces=ns)
                if ind is None:
                    continue
                if attr("firstLine") in ind.attrib:
                    document_first_line_count += 1
                if ind.get(attr("firstLineChars")) == "200":
                    document_first_line_chars_200_count += 1
            for table in document_root.xpath("//w:tbl", namespaces=ns):
                table_count += 1
                if table.xpath("./w:tr[1]/w:trPr/w:tblHeader", namespaces=ns):
                    tables_with_repeating_header += 1
                borders = table.find("w:tblPr/w:tblBorders", namespaces=ns)
                if borders is not None:
                    top = borders.find("w:top", namespaces=ns)
                    bottom = borders.find("w:bottom", namespaces=ns)
                    inside_h = borders.find("w:insideH", namespaces=ns)
                    inside_v = borders.find("w:insideV", namespaces=ns)
                    if (
                        top is not None
                        and bottom is not None
                        and top.get(attr("val")) == "single"
                        and bottom.get(attr("val")) == "single"
                        and (inside_h is None or inside_h.get(attr("val")) == "nil")
                        and (inside_v is None or inside_v.get(attr("val")) == "nil")
                    ):
                        tables_with_three_line_borders += 1
        footers = []
        for name in names:
            if name.startswith("word/footer") and name.endswith(".xml"):
                xml = archive.read(name).decode("utf-8")
                footers.append(
                    {
                        "part": name,
                        "left": 'w:jc w:val="left"' in xml,
                        "right": 'w:jc w:val="right"' in xml,
                        "center": 'w:jc w:val="center"' in xml,
                        "has_page_field": "PAGE" in xml,
                    }
                )
    return {
        "even_and_odd_headers": "w:evenAndOddHeaders" in settings_xml,
        "styles_first_line_count": styles_first_line_count,
        "document_first_line_count": document_first_line_count,
        "styles_first_line_chars_200_count": styles_first_line_chars_200_count,
        "document_first_line_chars_200_count": document_first_line_chars_200_count,
        "table_count": table_count,
        "tables_with_repeating_header": tables_with_repeating_header,
        "tables_with_three_line_borders": tables_with_three_line_borders,
        "footers": footers,
    }


def _validate_page_setup(first_section_mm: dict, errors: list[str], expected_page_setup_mm: dict | None = None) -> None:
    expected_page_setup_mm = expected_page_setup_mm or EXPECTED_PAGE_SETUP_MM
    for key, expected in expected_page_setup_mm.items():
        if expected is None:
            continue
        actual = first_section_mm.get(key)
        if actual is None:
            errors.append(f"Missing page setup value: {key}.")
            continue
        if abs(actual - expected) > PAGE_TOLERANCE_MM:
            errors.append(f"Page setup mismatch: {key} is {actual} mm, expected {expected} mm.")


def _validate_xml_diagnostics(xml: dict, errors: list[str], enabled_checks: set[str] | None = None) -> None:
    enabled_checks = enabled_checks or {
        "validation.indent_chars",
        "validation.footer_page_number",
        "validation.table_rules",
    }
    if "validation.footer_page_number" in enabled_checks and not xml.get("even_and_odd_headers"):
        errors.append("Odd/even footer setting is disabled: evenAndOddHeaders is missing.")
    if "validation.indent_chars" in enabled_checks and (xml.get("styles_first_line_count", 0) or xml.get("document_first_line_count", 0)):
        errors.append(
            "Point/cm first-line indentation detected: use w:firstLineChars instead of w:firstLine."
        )
    footers = xml.get("footers", [])
    if "validation.footer_page_number" in enabled_checks and footers:
        if any(footer.get("center") for footer in footers):
            errors.append("Centered page-number footer detected; official pages require odd right / even left.")
        if not any(footer.get("right") for footer in footers):
            errors.append("Missing right-aligned odd-page footer.")
        if not any(footer.get("left") for footer in footers):
            errors.append("Missing left-aligned even-page footer.")
    table_count = xml.get("table_count", 0)
    if "validation.table_rules" in enabled_checks and table_count:
        if xml.get("tables_with_repeating_header", 0) != table_count:
            errors.append("Some tables do not repeat their header row across page breaks.")
        if xml.get("tables_with_three_line_borders", 0) != table_count:
            errors.append("Some tables are not formatted as three-line tables.")


def _profile_validation_settings(profile_info: dict | None) -> tuple[list[str], dict, set[str]]:
    if not profile_info:
        return REQUIRED_STYLES, EXPECTED_PAGE_SETUP_MM, {
            "validation.required_styles",
            "validation.page_setup",
            "validation.indent_chars",
            "validation.footer_page_number",
            "validation.table_rules",
        }

    validation_rules = profile_info.get("validation_rules", {})
    required_styles = validation_rules.get("required_styles") or REQUIRED_STYLES
    page_setup = validation_rules.get("page_setup") or {}
    expected_page_setup = {
        "page_width": page_setup.get("page_width_mm"),
        "page_height": page_setup.get("page_height_mm"),
        "top_margin": page_setup.get("top_margin_mm"),
        "bottom_margin": page_setup.get("bottom_margin_mm"),
        "left_margin": page_setup.get("left_margin_mm"),
        "right_margin": page_setup.get("right_margin_mm"),
        "header_distance": page_setup.get("header_distance_mm"),
        "footer_distance": page_setup.get("footer_distance_mm"),
    }
    if not any(value is not None for value in expected_page_setup.values()):
        expected_page_setup = EXPECTED_PAGE_SETUP_MM
    enabled_checks = set(validation_rules.get("required_checks") or [])
    return required_styles, expected_page_setup, enabled_checks


def validate_document(
    docx_path: Path,
    output_path: Path | None = None,
    content_path: Path | None = None,
    profile_id: str | None = None,
    profile_info: dict | None = None,
) -> dict:
    warnings: list[str] = []
    errors: list[str] = []

    if profile_id and profile_info is None:
        try:
            profile_info = resolve_profile(profile_id)
        except ProfileError as exc:
            profile_info = {
                "profile_id": profile_id,
                "status": "invalid",
                "display_name": profile_id,
                "allow_formal_delivery": False,
                "feature_summary": {"implemented": [], "partial": [], "planned": [], "unknown": []},
            }
            errors.append(str(exc))

    required_styles, expected_page_setup, enabled_checks = _profile_validation_settings(profile_info)

    if profile_info and profile_info.get("status") != "production":
        warnings.append(
            f"Profile '{profile_info.get('profile_id')}' status is {profile_info.get('status')}; formal delivery is not allowed."
        )
    if profile_info and "validation.profile_rules" in enabled_checks:
        warnings.append("Profile-specific validation engine is partial/planned; report includes base checks only.")

    try:
        doc = Document(str(docx_path))
    except Exception as exc:  # pragma: no cover - user-facing failure path
        result = {
            "docx": str(docx_path),
            "ok": False,
            "errors": [f"Cannot open docx: {exc}"],
            "warnings": warnings,
        }
        if output_path:
            output_path.write_text(render_report(result), encoding="utf-8")
        return result

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    styles = sorted({p.style.name for p in paragraphs})
    all_style_names = {style.name for style in doc.styles}
    present_required = [name for name in required_styles if name in all_style_names]
    missing_required = [name for name in required_styles if name not in all_style_names]

    if not paragraphs:
        errors.append("Document has no non-empty paragraphs.")
    if "公文标题" not in styles and "Heading 1" not in styles:
        errors.append("No document-title style usage detected.")
    if "validation.required_styles" in enabled_checks and missing_required:
        errors.append("Missing preferred styles: " + ", ".join(missing_required))

    expected_title = None
    if content_path and content_path.exists():
        for line in content_path.read_text(encoding="utf-8-sig").splitlines():
            if line.startswith("# "):
                expected_title = line[2:].strip()
                break
    if expected_title and not any(expected_title in p.text for p in paragraphs):
        errors.append(f"Expected Markdown title not found in output: {expected_title}")

    section = doc.sections[0]
    first_section_mm = {
        "page_width": _mm(section.page_width),
        "page_height": _mm(section.page_height),
        "top_margin": _mm(section.top_margin),
        "bottom_margin": _mm(section.bottom_margin),
        "left_margin": _mm(section.left_margin),
        "right_margin": _mm(section.right_margin),
        "header_distance": _mm(section.header_distance),
        "footer_distance": _mm(section.footer_distance),
    }
    xml_diagnostics = _xml_diagnostics(docx_path)
    if "validation.page_setup" in enabled_checks:
        _validate_page_setup(first_section_mm, errors, expected_page_setup)
    _validate_xml_diagnostics(xml_diagnostics, errors, enabled_checks)

    result = {
        "docx": str(docx_path),
        "ok": not errors,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "styles_used": styles,
        "required_styles_present": present_required,
        "required_styles_missing": missing_required,
        "first_section_mm": first_section_mm,
        "xml_diagnostics": xml_diagnostics,
        "errors": errors,
        "warnings": warnings,
    }
    if profile_info:
        result["profile"] = {
            "profile_id": profile_info.get("profile_id"),
            "display_name": profile_info.get("display_name"),
            "status": profile_info.get("status"),
            "allow_formal_delivery": profile_info.get("allow_formal_delivery"),
            "template_path": profile_info.get("template_path"),
            "validation_rules_path": profile_info.get("validation_rules_path"),
            "feature_summary": profile_info.get("feature_summary", {}),
        }

    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(render_report(result), encoding="utf-8")
    return result


def render_report(result: dict) -> str:
    lines = [
        "# Validation Report",
        "",
        f"- Document: `{result.get('docx')}`",
        f"- OK: `{result.get('ok')}`",
        f"- Paragraphs: `{result.get('paragraph_count', 0)}`",
        f"- Tables: `{result.get('table_count', 0)}`",
        "",
    ]
    profile = result.get("profile")
    if profile:
        lines.extend(
            [
                "## Profile",
                "",
                f"- Profile ID: `{profile.get('profile_id')}`",
                f"- Display Name: `{profile.get('display_name')}`",
                f"- Status: `{profile.get('status')}`",
                f"- Allow Formal Delivery: `{profile.get('allow_formal_delivery')}`",
                f"- Template: `{profile.get('template_path')}`",
                f"- Validation Rules: `{profile.get('validation_rules_path')}`",
                "",
                "### Feature Status",
                "",
            ]
        )
        feature_summary = profile.get("feature_summary", {})
        for status in ("implemented", "partial", "planned", "unknown"):
            values = feature_summary.get(status, [])
            lines.append(f"- {status}: " + (", ".join(f"`{item}`" for item in values) if values else "none"))
        lines.append("")

        if profile.get("status") in {"draft", "planned"}:
            lines.extend(
                [
                    "### Delivery Gate",
                    "",
                    "- This profile is not production-ready.",
                    "- The document may be generated for testing only.",
                    "- Formal delivery requires a production profile and profile-specific validation.",
                    "",
                ]
            )

    lines.extend(["## Styles Used", ""])
    for style in result.get("styles_used", []):
        lines.append(f"- `{style}`")

    lines.extend(["", "## Required Styles", ""])
    present = result.get("required_styles_present", [])
    missing = result.get("required_styles_missing", [])
    lines.append("- Present: " + (", ".join(f"`{name}`" for name in present) if present else "none"))
    lines.append("- Missing: " + (", ".join(f"`{name}`" for name in missing) if missing else "none"))

    if result.get("first_section_mm"):
        lines.extend(["", "## First Section Page Setup", ""])
        for key, value in result["first_section_mm"].items():
            lines.append(f"- `{key}`: `{value}` mm")

    if result.get("xml_diagnostics"):
        xml = result["xml_diagnostics"]
        lines.extend(["", "## XML Diagnostics", ""])
        lines.append(f"- `even_and_odd_headers`: `{xml.get('even_and_odd_headers')}`")
        lines.append(f"- `styles_first_line_count`: `{xml.get('styles_first_line_count')}`")
        lines.append(f"- `document_first_line_count`: `{xml.get('document_first_line_count')}`")
        lines.append(f"- `styles_first_line_chars_200_count`: `{xml.get('styles_first_line_chars_200_count')}`")
        lines.append(f"- `document_first_line_chars_200_count`: `{xml.get('document_first_line_chars_200_count')}`")
        lines.append(f"- `table_count`: `{xml.get('table_count')}`")
        lines.append(f"- `tables_with_repeating_header`: `{xml.get('tables_with_repeating_header')}`")
        lines.append(f"- `tables_with_three_line_borders`: `{xml.get('tables_with_three_line_borders')}`")
        for footer in xml.get("footers", []):
            lines.append(
                "- Footer "
                f"`{footer.get('part')}`: "
                f"left=`{footer.get('left')}`, "
                f"right=`{footer.get('right')}`, "
                f"center=`{footer.get('center')}`, "
                f"has_page_field=`{footer.get('has_page_field')}`"
            )

    lines.extend(["", "## Errors", ""])
    errors = result.get("errors", [])
    lines.extend([f"- {item}" for item in errors] if errors else ["- none"])

    lines.extend(["", "## Warnings", ""])
    warnings = result.get("warnings", [])
    lines.extend([f"- {item}" for item in warnings] if warnings else ["- none"])
    lines.append("")
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="Validate a generated Word document.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--content", type=Path)
    parser.add_argument("--profile", help="Optional document profile id for profile-aware validation metadata and rules.")
    args = parser.parse_args()
    validate_document(args.docx, args.output, args.content, profile_id=args.profile)


if __name__ == "__main__":
    main()
